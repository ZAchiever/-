from docx import Document
# document = Document('storage/合同.docx')
import datetime
import os
# DICT = {
#     "##托运公司填写处##": "new",
#     '##地址##': '',
#  datetime.now().strftime( '%d_%H_%M_%S')
# }

import zipfile


# history = pd.read_excel('history.xlsx', sheet_name=0, index_col=0)
# history.to_excel('storage\\car_2.xlsx')

def zip_file(src_dir):
    zip_name = src_dir + '.zip'
    z = zipfile.ZipFile(zip_name, 'w', zipfile.ZIP_DEFLATED)
    for dirpath, dirnames, filenames in os.walk(src_dir):
        fpath = dirpath.replace(src_dir, '')
        fpath = fpath and fpath + os.sep or ''
        for filename in filenames:
            z.write(os.path.join(dirpath, filename), fpath+filename)
            print('==压缩成功==')
    z.close()


def mkdir(path):

    folder = os.path.exists(path)

    if not folder:  # 判断是否存在文件夹如果不存在则创建为文件夹
        os.makedirs(path)  # makedirs 创建文件时如果路径不存在会创建这个路
    else:
        print('文件夹已存在')


def check(DICT, doc_name):
    # tables
    document = Document('storage//'+doc_name)

    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                for key, value in DICT.items():
                    if key in table.cell(row, col).text:
                        print(key+"->"+value)
                        table.cell(row, col).text = table.cell(
                            row, col).text.replace(key, value)

    # paragraphs
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in DICT.items():
                if key in para.runs[i].text:
                    print(key+"->"+value)
                    para.runs[i].text = para.runs[i].text.replace(key, value)

    return document


def out_put_doc(DICT):
    doc_list = ['合同.docx', '提货单.docx', '托运单.docx']
    hash_dir = 'all_you_need//'+datetime.datetime.now().strftime('%d_%H_%M') + \
        DICT['##托运公司##']
    for i in doc_list:
        doc = check(DICT, i)
        mkdir(hash_dir)
        doc.save(hash_dir+'//'+i)
    zip_file(hash_dir)


# print(check(document))
# document.save('word1.docx')
